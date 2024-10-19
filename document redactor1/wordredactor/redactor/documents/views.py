from django.shortcuts import render, redirect, get_object_or_404
from .forms import DocumentForm
from .models import Document  
from docx import Document as DocxDocument
from docx.shared import RGBColor
from io import BytesIO
import base64
from PIL import Image
from bs4 import BeautifulSoup

def upload_document(request):
    documents = Document.objects.all()
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('upload_document')
    else:
        form = DocumentForm()
    return render(request, 'documents/upload.html', {'form': form, 'documents': documents})

def edit_document(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    docx_content = ""
    
    if document.file.name.endswith('.docx'):
        docx_path = document.file.path
        doc = DocxDocument(docx_path)

        
        for para in doc.paragraphs:
            for run in para.runs:
                if run._element.xpath('.//w:drawing'):
                    for drawing in run._element.xpath('.//a:blip'):
                        img_part = doc.part.related_parts[drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')]
                        img_data = img_part.blob
                        img_base64 = base64.b64encode(img_data).decode('utf-8')
                        docx_content += f'<img src="data:image/png;base64,{img_base64}" />'
                docx_content += run.text + "<br/>"

    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES, instance=document)
        new_content = request.POST.get('docx_content', '')
        if form.is_valid():
            form.save()

            if document.file.name.endswith('.docx') and new_content:
                doc = DocxDocument()
                soup = BeautifulSoup(new_content, 'html.parser')

                for element in soup.find_all(['p', 'div', 'span', 'img', 'ol', 'ul']):
                    doc = html_to_docx(element, doc)

                doc.save(docx_path)

            return redirect('upload_document')
    else:
        form = DocumentForm(instance=document)

    return render(request, 'documents/edit.html', {
        'form': form,
        'document': document,
        'docx_content': docx_content
    })

def html_to_docx(element, doc):
    
    if element.name == 'img':
       
        img_data = element['src'].split(',')[1]
        img = Image.open(BytesIO(base64.b64decode(img_data)))
        image_stream = BytesIO()
        img.save(image_stream, format='PNG')
        image_stream.seek(0)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_stream)
    elif element.name == 'ol':
        for li in element.find_all('li'):
            list_item = doc.add_paragraph(style='ListNumber')
            list_item.add_run(li.get_text())
    elif element.name == 'ul':
        for li in element.find_all('li'):
            list_item = doc.add_paragraph(style='ListBullet')
            list_item.add_run(li.get_text())
    else:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(element.get_text())

        
        if 'style' in element.attrs:
            styles = element['style']
            if 'color:' in styles:
                color = styles.split('color:')[1].split(';')[0].strip()
                if color.startswith('#'):
                    rgb_color = color[1:]
                    run.font.color.rgb = RGBColor(
                        int(rgb_color[0:2], 16), 
                        int(rgb_color[2:4], 16), 
                        int(rgb_color[4:6], 16)
                    )
            if 'background-color:' in styles:
               
                pass

    
        if element.name == 'strong' or 'bold' in element.attrs.get('class', []):
            run.bold = True
        if element.name == 'em' or 'italic' in element.attrs.get('class', []):
            run.italic = True

    return doc
